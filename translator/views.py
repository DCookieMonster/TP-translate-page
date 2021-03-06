from django.shortcuts import render,render_to_response,RequestContext
import os
from .form import DocumentForm
from .models import Paper
from .models import Paragraph
from .models import Translated_Paragraph
import datetime
from django.utils.encoding import smart_unicode
from django.conf import settings
from django.contrib.auth import authenticate, login
from docx import Document
import random
import string
from random import randint
import Docx
from cStringIO import StringIO
from docx import Document
from django.http import HttpResponse
from django.contrib.auth.models import User
from django.contrib.auth.decorators import login_required
from django.contrib.auth import logout
import datetime
import RegexTxt
import urllib2
import googleTranslator
# Create your views here.


def home(request):
    if(request.method=='POST'):
            if(not Paper.objects.filter(code=request.POST['code'])):
                context={
               'error':'<div class="alert alert-danger alert-dismissible" role="alert">'
                            ' <button type="button" class="close" data-dismiss="alert" aria-label="Close"><span aria-hidden="true">&times;</span></button>'
                            'Code Not Found </div>',
                 }
                return render(request,"index.html",context)

            para=GetParagraph(request.POST['code'])

            #return paragparh to translate
            #https://translate.google.co.il/#en/iw/dsadsgasdgasdg
            tmpTxt=RegexTxt.RepleceTxt(para.txt)
            link='https://translate.google.co.il/#en/iw/'+tmpTxt
            context={
                'paraNum': para.num,
                'txt': para.txt,
                'id': para.id,
                'link': link,
            }
            return render(request,"translate.html",context)


    return render_to_response("index.html",locals(),context_instance=RequestContext(request))


def GetParagraph(code):


    paper=Paper.objects.filter(code=code)[0]
    if(check_if_ready(paper)):
        #the doc is ready
        print("Ready")
        Docx.WriteDocx(paper)

    paragprahs=Paragraph.objects.filter(paperId=paper.id)
    num=randint(0,paragprahs.__len__()-1)
    if(not Translated_Paragraph.objects.filter(paraId=paragprahs[num]).exists()):
        return paragprahs[num]
    for para in paragprahs:
       if(Translated_Paragraph.objects.filter(paraId=para.id).exists()):
           continue
       return para


def register(request):

    if request.POST.get('login-submit'):
        users = User.objects.all()
        try:
            user = authenticate(username=request.POST.get('username'),password=request.POST.get('password'))
            # user=users.filter(username=request.POST.get('username'),password=request.POST.get('password'))[0]
            # if (user ):
            if user is not None:
                #approve user
                if user.is_active:
                    login(request, user)
                response=render_to_response("index.html",locals(),context_instance=RequestContext(request))
                set_cookie(response,'userId',user.id,days_expire=1)
                return response
        except:
            #dont approve user
            context={
               'error':'<div class="alert alert-danger alert-dismissible" role="alert">'
                            ' <button type="button" class="close" data-dismiss="alert" aria-label="Close"><span aria-hidden="true">&times;</span></button>'
                            'Wrong Username or Password </div>',
            }
            return render(request,"register.html",context)




    if(request.POST.get('register-submit')):
        if(request.POST.get('password')!=request.POST.get('confirm-password')):
            context={
                'error':'<div class="alert alert-warning alert-dismissible" role="alert">'
                        ' <button type="button" class="close" data-dismiss="alert" aria-label="Close"><span aria-hidden="true">&times;</span></button>'
                        'Password don\'t match </div>',
            }
            return render_to_response("register.html",locals(),context)
        if User.objects.filter(username=request.POST['username']).exists():
            context={
                'error':'<div class="alert alert-warning alert-dismissible" role="alert">'
                        ' <button type="button" class="close" data-dismiss="alert" aria-label="Close"><span aria-hidden="true">&times;</span></button>'
                        'Username already exists in the system </div>',
            }
            return render(request,'register.html',context)
        user = User.objects.create_user(smart_unicode(request.POST.get('username')), smart_unicode(request.POST.get('email')), smart_unicode(request.POST.get('password')))
        # user = User()
        # user.username=smart_unicode(request.POST.get('username'))
        # user.password=smart_unicode(request.POST.get('password'))
        # user.email=smart_unicode(request.POST.get('email'))
        user.first_name=smart_unicode(request.POST.get('fname'))
        user.last_name=smart_unicode(request.POST.get('lname'))
        # user.timestamp=smart_unicode(datetime.datetime.today())
        user.save()
        user = authenticate(username=request.POST.get('username'),password=request.POST.get('password'))
        login(request, user)
        context={
            'title': "Thank You for Registered!",
            'subTitle': "You Can Start Uploading",
        }
        response=render_to_response("index.html",locals(),context_instance=RequestContext(request))
        set_cookie(response, 'userId', user.id,days_expire=1)
        return response
    return render(request, "register.html",{})




#set cookie after sign in or register
def set_cookie(response, key, value, days_expire = 7):
  if days_expire is None:
    max_age = 365 * 24 * 60 * 60  #one year
  else:
    max_age = days_expire * 24 * 60 * 60
  expires = datetime.datetime.strftime(datetime.datetime.utcnow() + datetime.timedelta(seconds=max_age), "%a, %d-%b-%Y %H:%M:%S GMT")
  response.set_cookie(key, value, max_age=max_age, expires=expires, domain=settings.SESSION_COOKIE_DOMAIN, secure=settings.SESSION_COOKIE_SECURE or None)

def id_generator(size=6, chars=string.ascii_uppercase + string.digits):
    code= ''.join(random.choice(chars) for _ in range(size)).lower()
    while(Paper.objects.filter(code=code).exists()):
        code= ''.join(random.choice(chars) for _ in range(size)).lower()
    return code



@login_required
def uploadFile(request):
    # Handle file upload


    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            document = Document(request.FILES['docfile'])
            code = id_generator()
            paper = Paper()
            paper.code = code
            paper.name = code+"-"+request.FILES['docfile']._name
            paper.userId = User.objects.filter(id=request.user.id)[0]
            paper.docx = request.FILES['docfile']
            paper.save()
            # Divde to paragraph
            par = ""
            count = 0
            for paragraph in document.paragraphs:
                par +=paragraph.text.encode('utf8')
                if par.__len__()<30:
                    par += '<br>'
                    continue
                insertParagraphToDB(par,count,paper.id)
                count += 1
                par = ""
            # Redirect to the document list after POST
            context = {
                'code': code,
            }
            return render(request, 'getCode.html', context)
    else:
        form = DocumentForm() # A empty, unbound formxw

    # Render list page with the documents and the form
    return render_to_response('upload.html', locals(), context_instance=RequestContext(request))


def insertParagraphToDB(paragraph,paraNum,paperId):
    para=Paragraph()
    para.num=paraNum
    para.txt=paragraph
    para.paperId=Paper.objects.filter(id=paperId)[0]
    para.save()


def translate(request):
    if request.method == 'POST':
        para = Paragraph.objects.filter(id=request.POST['id'])[0]
        trans = Translated_Paragraph()
        trans.txt = request.POST['txt-tran']
        trans.paraId = para
        trans.save()
        paper = Paper.objects.filter(id=para.paperId.id)[0]
        context = {
            'code': paper.code,
        }
        if(check_if_ready(paper)):
            Docx.WriteDocx(paper)

            #write to file

        return render(request, 'thankyou.html', context)
    else:
        return render_to_response('index.html', locals(), context_instance=RequestContext(request))


def check_if_ready(paper):
    paragraphs=Paragraph.objects.filter(paperId=paper)
    for para in paragraphs:
        if(not Translated_Paragraph.objects.filter(paraId=para).exists()):
            return False
    return True


# get the original File
def GetFile(request):
    if(request.method=='POST'):

        para=Paragraph.objects.filter(id=request.POST['paraId'])[0]
        paper=Paper.objects.filter(id=para.paperId.id)[0]
        document = Document(paper.docx)
        f = StringIO()
        document.save(f)
        length = f.tell()
        f.seek(0)
        response = HttpResponse(
            f.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename='+paper.name
        response['Content-Length'] = length
        return response

@login_required
def logout_view(request):
    logout(request)
    return render_to_response('index.html',locals(),context_instance=RequestContext(request))


@login_required
def profile_page(request):
    user = User.objects.filter(id=request.user.id)[0]
    papers = Paper.objects.filter(userId=user)
    statusList = []
    for paper in papers:
        paragraphs = Paragraph.objects.filter(paperId=paper)
        count = 0.0
        for paragraph in paragraphs:
            if Translated_Paragraph.objects.filter(paraId=paragraph).exists():
                count += 1.0
        avg = count/paragraphs.__len__()
        statusList.append((paper.id,paper.name,paper.code,avg*100))

    context = {
        'status': statusList,
    }
    return render(request,'profilepage.html',context)

# get the original File
def get_translated_file(request):
    if request.method == 'POST':

        para=Paragraph.objects.filter(id=request.POST['paraId'])[0]
        paper=Paper.objects.filter(id=para.paperId.id)[0]
        document = Docx.WriteDocx(paper)
        f = StringIO()
        document.save(f)
        length = f.tell()
        f.seek(0)
        response = HttpResponse(
            f.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=translated-'+paper.name
        response['Content-Length'] = length
        return response

